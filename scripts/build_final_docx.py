#!/usr/bin/env python3
"""
output/json/{pdf1,pdf2}/*.json をページ順に集約して .docx を生成する。

依存：
    pip install python-docx

実行：
    python3 scripts/build_final_docx.py
    python3 scripts/build_final_docx.py --pdf pdf1     # 片方だけ
    python3 scripts/build_final_docx.py --include-uncertain    # 不確実箇所を末尾に付録として出す
"""

from __future__ import annotations
import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
JSON_ROOT = ROOT / "output" / "json"
OUT_DIR = ROOT / "output" / "docx"

CITATION_OPEN = re.compile(r"〔引用:[^〕]*〕")
CITATION_CLOSE = re.compile(r"〔引用ここまで〕")


def load_pages(pdf_key: str) -> list[dict]:
    pdf_dir = JSON_ROOT / pdf_key
    if not pdf_dir.exists():
        return []
    files = sorted(pdf_dir.glob("page-*.json"))
    pages: list[dict] = []
    for f in files:
        try:
            pages.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception as e:
            print(f"[WARN] {f.name}: {e}")
    return pages


def add_styled_body(doc: Document, body: str) -> None:
    """body文字列を段落単位で追加し、〔引用:..〕〜〔引用ここまで〕を引用ブロックとして
    視覚的にインデント付きで描画する。"""
    if not body:
        return
    in_quote = False
    for raw_para in body.split("\n"):
        text = raw_para.rstrip()
        if not text:
            doc.add_paragraph("")
            continue

        # 引用開始
        if CITATION_OPEN.search(text):
            in_quote = True
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
            run.font.size = Pt(10.5)
            continue

        # 引用終了
        if CITATION_CLOSE.search(text):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
            run.font.size = Pt(10.5)
            in_quote = False
            continue

        p = doc.add_paragraph()
        if in_quote:
            p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        if in_quote:
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)


def add_page(doc: Document, page: dict, pdf_label: str) -> None:
    h = doc.add_heading(
        f"[{pdf_label}] p.{page.get('page_number')} ({page.get('page_label', '')})  {page.get('section', '')}".strip(),
        level=2,
    )
    body = page.get("body", "") or ""
    add_styled_body(doc, body)

    for tbl in page.get("tables", []) or []:
        cap = tbl.get("caption") or tbl.get("id", "table")
        doc.add_paragraph(f"[表] {cap}").runs[0].bold = True
        for row_line in (tbl.get("markdown") or "").split("\n"):
            doc.add_paragraph(row_line)

    fns = page.get("footnotes", []) or []
    if fns:
        p = doc.add_paragraph()
        p.add_run("[欄外・注記]").bold = True
        for fn in fns:
            doc.add_paragraph(
                f"  ・({fn.get('type','')}) {fn.get('anchor','')}: {fn.get('text','')}"
            )


def add_uncertain_appendix(doc: Document, all_pages: list[tuple[str, dict]]) -> None:
    doc.add_page_break()
    doc.add_heading("付録：自信のない文字一覧", level=1)
    n = 0
    for pdf_label, p in all_pages:
        for u in p.get("uncertain", []) or []:
            n += 1
            doc.add_paragraph(
                f"[{pdf_label} p.{p.get('page_number')}] {u.get('location','')} "
                f"→ 認識: 「{u.get('original_recognition','')}」 / {u.get('note','')}"
            )
    if n == 0:
        doc.add_paragraph("（該当なし）")
    else:
        doc.add_paragraph(f"（合計 {n} 件）")


def build(pdf_keys: list[str], include_uncertain: bool) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "MS Mincho"
    style.font.size = Pt(10.5)

    doc.add_heading("高圧ガス保安法規集（第23次改訂版）OCR成果", level=0)
    doc.add_paragraph(
        "本書は『高圧ガス保安法規集 第23次改訂版』（特別民間法人 高圧ガス保安協会 編）"
        "のスキャン画像を Claude Code でOCR・構造化した社内RAG用テキスト化成果である。"
        "赤色で示した本文は色刷り引用条文ブロックを表す。"
    )

    flat: list[tuple[str, dict]] = []
    for key in pdf_keys:
        pages = load_pages(key)
        if not pages:
            continue
        doc.add_page_break()
        doc.add_heading(
            f"{'第一巻' if key == 'pdf1' else '第二巻'} ({key})  全 {len(pages)} ページ",
            level=1,
        )
        for p in pages:
            add_page(doc, p, key)
            flat.append((key, p))

    if include_uncertain:
        add_uncertain_appendix(doc, flat)

    out_path = OUT_DIR / "koatsu_gas_houkishuu_v23.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--pdf",
        choices=["pdf1", "pdf2", "all"],
        default="all",
        help="集約対象（既定: all = 両方）",
    )
    ap.add_argument(
        "--include-uncertain",
        action="store_true",
        help="末尾に uncertain の一覧を付録として追加",
    )
    args = ap.parse_args()

    pdf_keys = ["pdf1", "pdf2"] if args.pdf == "all" else [args.pdf]
    out = build(pdf_keys, args.include_uncertain)
    print(f"[OK] generated: {out.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
