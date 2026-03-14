import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> list[dict]:
    """Extract text from PDF file, returning text with page numbers."""
    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                pages.append({"text": text.strip(), "page": i})
    return pages


def extract_text_from_docx(file_bytes: bytes) -> list[dict]:
    """Extract text from Word document."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    if full_text:
        return [{"text": full_text, "page": None}]
    return []


def extract_text(file_bytes: bytes, filename: str) -> list[dict]:
    """Extract text from a file based on its extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"サポートされていないファイル形式です: {filename}")
